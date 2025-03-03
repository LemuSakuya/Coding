from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_invitation():
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "缘聚十年情定终生"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = header_para.runs[0]
    run.font.size = Pt(24)
    run.font.color.rgb = parse_xml(r'<a:schemeClr val="accent1"/>')

    p1 = doc.add_paragraph("尊敬的")
    p1.add_run("《姓名》").bold = True
    p1.add_run("先生/女士：")
    p1.style = "要点"
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(12)

    body_text = """十年前的夏天，我们满怀憧憬告别母校...（正文内容同上）"""
    p2 = doc.add_paragraph(body_text)
    p2.style = "正文"
    p2.paragraph_format.first_line_indent = Inches(0.2)

    doc.add_paragraph("附：毕业十周年聚会日程安排表").style = "要点"
    table = doc.add_table(rows=11, cols=4)
    table.style = "Medium Shading 1 Accent 4"

    for row in table.rows:
        row.height = Cm(0.7)

    data_rows = [
        ["10月1日", "16:00", "报到，入住酒店", "新世纪大酒店"],
        # ... 补充其余行数据
    ]
    for i, row_data in enumerate(data_rows):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = cell_data

    p_date = doc.add_paragraph("2014年9月20日")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(12)
    p_date.runs[0].bold = True

    doc.save("邀请函.docx")

def create_classmate_list():
    doc = Document()
    table = doc.add_table(rows=9, cols=4)
    table.style = "Table Grid"

    headers = ["编号", "姓名", "性别", "联系方式"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    data = [
        ["1", "张珮珮", "女", "138-XXXX-1234"],
        # ... 补充其他同学数据
    ]
    for i, row_data in enumerate(data, start=1):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = cell_data
    doc.save("同学录.docx")

if __name__ == "__main__":
    create_invitation()
    create_classmate_list()
